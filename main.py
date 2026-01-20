import os
import re
import io
import argparse
import tempfile
from pathlib import Path
from typing import Optional

from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from urllib.parse import quote
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse


def clean_llm_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"^\s*(---|\*\*\*|___)\s*$", "", text, flags=re.M)
    text = text.replace("**", "")
    text = re.sub(r"^\s*#{1,6}\s*", "", text, flags=re.M)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    text = "\n".join([ln.rstrip() for ln in text.split("\n")])
    return text.strip()


def force_heading_breaks(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"([。！？；])\s*([一二三四五六七八九十]+、)", r"\1\n\n\2", text)
    text = re.sub(r"([。！？；])\s*([（(][一二三四五六七八九十]+[)）])", r"\1\n\n\2", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    return text.strip()


def extract_title_and_body(raw_llm_text: str):
    cleaned = clean_llm_text(raw_llm_text)
    lines = [ln.strip() for ln in cleaned.split("\n")]
    title_keywords = ["总体统筹方案", "统筹方案", "总体方案", "工作方案", "实施方案", "方案"]
    candidates = []
    for idx, ln in enumerate(lines[:15]):
        if not ln:
            continue
        if re.match(r"^[一二三四五六七八九十]+、", ln):
            continue
        if re.match(r"^[（(][一二三四五六七八九十]+[)）]", ln):
            continue
        score = 0
        for kw in title_keywords:
            if kw in ln:
                score += 10
        if 8 <= len(ln) <= 60:
            score += 3
        if "公司" in ln:
            score += 2
        candidates.append((score, idx, ln))
    if candidates:
        candidates.sort(reverse=True)
        _, title_idx, title = candidates[0]
    else:
        title_idx, title = next(((i, ln) for i, ln in enumerate(lines) if ln), (0, ""))
    body_lines = [ln for i, ln in enumerate(lines) if i != title_idx]
    body = "\n".join(body_lines).strip()
    body = clean_llm_text(body)
    return title.strip(), body


def set_run_font(run, name, size_pt=None, bold=None, color_rgb=(0, 0, 0)):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bool(bold)
    run.font.color.rgb = RGBColor(*color_rgb)


def apply_style_defaults(doc: Document, font_name="宋体", body_size=12, h1_size=16, h2_size=14, color_rgb=(0, 0, 0)):
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.font.size = Pt(body_size)
    normal.font.color.rgb = RGBColor(*color_rgb)

    h1 = doc.styles["Heading 1"]
    h1.font.name = font_name
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    h1.font.size = Pt(h1_size)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(*color_rgb)

    h2 = doc.styles["Heading 2"]
    h2.font.name = font_name
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    h2.font.size = Pt(h2_size)
    h2.font.bold = False
    h2.font.color.rgb = RGBColor(*color_rgb)


def add_main_title(doc: Document, title: str, font_name="宋体", title_size=20, color_rgb=(0, 0, 0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.strip())
    set_run_font(run, font_name, title_size, True, color_rgb)


def add_paragraph_with_style(doc: Document, text: str, style_name: str, font_name: str, size_pt: int, bold: bool, color_rgb=(0, 0, 0)):
    p = doc.add_paragraph("", style=style_name)
    run = p.add_run(text)
    set_run_font(run, font_name, size_pt, bold, color_rgb)
    return p


def write_llm_body_to_doc(doc: Document, body_text: str, font_name="宋体", body_size=12, h1_size=16, h2_size=14, color_rgb=(0, 0, 0)):
    for block in body_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if re.match(r"^[一二三四五六七八九十]+、", block):
            add_paragraph_with_style(doc, block, "Heading 1", font_name, h1_size, True, color_rgb)
            continue
        if re.match(r"^[（(][一二三四五六七八九十]+[)）]", block):
            add_paragraph_with_style(doc, block, "Heading 2", font_name, h2_size, False, color_rgb)
            continue
        block = block.replace("\n", "")
        add_paragraph_with_style(doc, block, "Normal", font_name, body_size, False, color_rgb)


def append_fixed_doc_as_text(doc_out: Document, fixed_doc: Document, font_name="宋体", body_size=12, h1_size=16, h2_size=14, color_rgb=(0, 0, 0)):
    for p in fixed_doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        if re.match(r"^[一二三四五六七八九十]+、", txt):
            add_paragraph_with_style(doc_out, txt, "Heading 1", font_name, h1_size, True, color_rgb)
        elif re.match(r"^[（(][一二三四五六七八九十]+[)）]", txt):
            add_paragraph_with_style(doc_out, txt, "Heading 2", font_name, h2_size, False, color_rgb)
        else:
            add_paragraph_with_style(doc_out, txt, "Normal", font_name, body_size, False, color_rgb)


def build_prompt(prompt_path: str | None) -> str:
    return """根据公司的组织架构思考并撰写统筹方案
## 工作步骤 
1. 首先生成一段概括性描述，用于阐述本方案的目的
2. 根据活动内容，分析公司参与该活动的目的，即第一章工作背景。

### 第一步 概括制作本方案的目的
用一段话说明撰写本方案的目的，请注意，这一部分不是一个章节，只是文档前的一个总结性部分。对于计划参与的时间以及具体活动内容名称等信息，应根据从输入文件中分析出。
例：
南方电网公司参加2025年世界人工智能大会总体统筹方案

为落实国家关于加快人工智能与实体经济深度融合、深入实施“AI+”行动的战略部署，展示公司“AI+”行动最新成果和实践经验，搭建开放、共享、合作的行业AI交流平台，公司计划于2025年7月下旬参加2025年世界人工智能大会（WAIC，以下简称“大会”），为组织做好会议筹备工作，特制定本统筹方案。
### 第二步 撰写一、工作背景
根据活动内容，分析说明公司参与该活动能起到什么样的宣传作用，其中关于活动的一些具体信息需要根据活动内容分析得出。"""


def generate_docx_bytes(
    activity_docx_bytes: bytes,
    fixed_docx_bytes: bytes,
    model: str,
    prompt: str,
    font_name: str,
    title_size: int,
    h1_size: int,
    h2_size: int,
    body_size: int,
) -> bytes:
    api_key = os.getenv("DASHSCOPE_API_KEY")
    if not api_key:
        raise RuntimeError("DASHSCOPE_API_KEY is not set")

    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    )

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tf:
            tf.write(activity_docx_bytes)
            tmp_path = tf.name

        file_object = client.files.create(
            file=Path(tmp_path),
            purpose="file-extract"
        )

        completion = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "你是南方电网公司策展宣传的文稿撰写专家，你的目标是制定南方电网公司参加输入活动的总体统筹方案。"},
                {"role": "system", "content": f"fileid://{file_object.id}"},
                {"role": "user", "content": prompt}
            ],
            stream=False
        )

        raw_llm_text = completion.choices[0].message.content or ""
        title, body = extract_title_and_body(raw_llm_text)
        if not title:
            title = "南方电网公司参加相关活动总体统筹方案"
        body = force_heading_breaks(body)

        fixed_doc = Document(io.BytesIO(fixed_docx_bytes))
        out_doc = Document()

        apply_style_defaults(
            out_doc,
            font_name=font_name,
            body_size=body_size,
            h1_size=h1_size,
            h2_size=h2_size,
            color_rgb=(0, 0, 0)
        )

        add_main_title(out_doc, title, font_name=font_name, title_size=title_size, color_rgb=(0, 0, 0))
        out_doc.add_paragraph("")
        write_llm_body_to_doc(out_doc, body, font_name=font_name, body_size=body_size, h1_size=h1_size, h2_size=h2_size, color_rgb=(0, 0, 0))
        out_doc.add_paragraph("")
        append_fixed_doc_as_text(out_doc, fixed_doc, font_name=font_name, body_size=body_size, h1_size=h1_size, h2_size=h2_size, color_rgb=(0, 0, 0))

        # buf = io.BytesIO()
        # out_doc.save(buf)
        buf = io.BytesIO()
        out_doc.save(buf)
        out_bytes = buf.getvalue()
        return out_bytes
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


app = FastAPI(title="Docx Plan Generator")


@app.post("/generate")
async def generate(
    activity_docx: UploadFile = File(...),
    model: str = Form("qwen-long"),
    prompt: Optional[str] = Form(None),
    font_name: str = Form("宋体"),
    title_size: int = Form(20),
    h1_size: int = Form(16),
    h2_size: int = Form(14),
    body_size: int = Form(12),
):
    if not activity_docx.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="activity_docx must be .docx")

    activity_bytes = await activity_docx.read()
    BASE_DIR = Path(__file__).resolve().parent
    FIXED_DOCX_PATH = BASE_DIR / "1.南方电网公司参加XXXXX总体统筹方案.docx"

    if not FIXED_DOCX_PATH.exists():
        raise HTTPException(status_code=500, detail="fixed_docx template not found")

    with open(FIXED_DOCX_PATH, "rb") as f:
        fixed_bytes = f.read()

    prompt_text = build_prompt(prompt)

    try:
        out_bytes = generate_docx_bytes(
            activity_docx_bytes=activity_bytes,
            fixed_docx_bytes=fixed_bytes,
            model=model,
            prompt=prompt_text,
            font_name=font_name,
            title_size=title_size,
            h1_size=h1_size,
            h2_size=h2_size,
            body_size=body_size,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    filename = "活动统筹方案.docx"
    headers = {
        "Content-Disposition": "attachment; filename=output.docx; filename*=UTF-8''" + quote(filename)
    }

    return StreamingResponse(
        io.BytesIO(out_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--host", default="0.0.0.0")
    parser.add_argument("--port", type=int, default=8000)
    args = parser.parse_args()
    import uvicorn
    uvicorn.run("main:app", host=args.host, port=args.port)


if __name__ == "__main__":
    main()
